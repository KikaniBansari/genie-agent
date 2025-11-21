import os
import json
from datetime import datetime
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, asdict
from dotenv import load_dotenv
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from langchain_groq import ChatGroq

load_dotenv()

logger = logging.getLogger(__name__)

# --- Configuration & Models ---
@dataclass
class ResearchResult:
    summary: str
    sources: List[Dict]
    confidence: float = 0.0

@dataclass
class StrategyStep:
    title: str
    description: str
    priority: str  # "high", "medium", "low"
    timeline: str

class GenieEngine:
    def __init__(self):
        self.groq_key = os.getenv("GROQ_API_KEY")
        self.tavily_key = os.getenv("TAVILY_API_KEY")
        
        # Initialize LangChain with Groq
        if self.groq_key:
            self.llm = ChatGroq(
                model="mixtral-8x7b-32768",
                groq_api_key=self.groq_key,
                temperature=0.7,
                max_tokens=2000
            )
        else:
            logger.warning("GROQ_API_KEY not found. Using fallback mode.")
            self.llm = None
    
    def _llm_generate(self, prompt: str, model: Optional[str] = None, temperature: float = 0.7) -> str:
        """Secure Server-Side LLM Call using LangChain with Groq"""
        
        if self.llm is None:
            return self._fallback_response(prompt)
        
        try:
            # Create a new LLM instance with the requested temperature
            llm = ChatGroq(
                model=model or "moonshotai/kimi-k2-instruct",
                groq_api_key=self.groq_key,
                temperature=temperature,
                max_tokens=2000
            )
            
            # Use LangChain to generate response
            response = llm.invoke(prompt)
            return response.content
        except Exception as e:
            logger.error(f"LangChain Groq API error: {str(e)}")
            return f"Error generating response: {str(e)}"
    
    def _fallback_response(self, prompt: str) -> str:
        """Fallback response when no API keys are available"""
        if "strategy" in prompt.lower():
            return """Based on your business goal, here's a recommended strategy:

1. **Market Research & Analysis**
   - Conduct thorough market research to understand your target audience
   - Analyze competitor strategies and identify gaps
   - Define your unique value proposition

2. **Strategic Planning**
   - Set clear, measurable objectives
   - Develop a detailed action plan with timelines
   - Allocate resources effectively

3. **Execution & Optimization**
   - Implement your strategy with regular monitoring
   - Gather feedback and iterate
   - Scale successful initiatives

Note: For AI-powered insights, please configure API keys in your environment variables."""
        return "I'm ready to help! Please configure your API keys to enable full AI capabilities."
    
    def search_web(self, query: str, max_results: int = 5) -> List[Dict]:
        """Secure Server-Side Search with enhanced error handling"""
        if not self.tavily_key:
            logger.warning("Tavily API key not found. Using fallback search.")
            return self._fallback_search(query)
        
        try:
            resp = requests.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": self.tavily_key,
                    "query": query,
                    "max_results": max_results,
                    "search_depth": "advanced"
                },
                timeout=15
            )
            resp.raise_for_status()
            results = resp.json().get("results", [])
            logger.info(f"Found {len(results)} search results for query: {query[:50]}")
            return results
        except requests.exceptions.Timeout:
            logger.error("Search request timed out")
            return self._fallback_search(query)
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            return self._fallback_search(query)
    
    def _fallback_search(self, query: str) -> List[Dict]:
        """Fallback search results when API is unavailable"""
        return [{
            "title": f"Research: {query}",
            "content": "Market research data is being processed. For real-time data, please configure the Tavily API key.",
            "url": "https://example.com",
            "score": 0.5
        }]
    
    def _extract_industry_insights(self, industry: str, search_results: List[Dict]) -> str:
        """Extract industry-specific insights from search results"""
        if not search_results:
            return f"General insights for {industry} industry."
        
        content = "\n".join([r.get('content', '')[:300] for r in search_results[:3]])
        return content[:1000]  # Limit context size
    
    def run_pipeline(self, goal: str, industry: str) -> Dict[str, Any]:
        """Enhanced Orchestrator with multiple phases"""
        
        logger.info(f"Starting pipeline for goal: {goal[:50]}... in industry: {industry}")
        
        # Phase 1: Enhanced Research
        research_queries = [
            f"{industry} market trends 2024",
            f"{industry} {goal} best practices",
            f"{industry} competitive landscape"
        ]
        
        all_search_results = []
        for query in research_queries:
            results = self.search_web(query, max_results=3)
            all_search_results.extend(results)
        
        # Deduplicate results
        seen_urls = set()
        unique_results = []
        for result in all_search_results:
            url = result.get('url', '')
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_results.append(result)
        
        context_text = "\n\n".join([r.get('content', '') for r in unique_results[:5]])
        industry_insights = self._extract_industry_insights(industry, unique_results)
        
        # Phase 2: Enhanced Strategy Generation
        strategy_prompt = f"""You are an expert business strategist. Based on the following context, create a comprehensive 5-step strategic plan.

Context and Market Research:
{context_text[:3000]}

Business Goal: {goal}
Industry: {industry}

Generate a detailed strategic plan with:
1. Executive Summary (2-3 sentences)
2. Five actionable strategic steps, each with:
   - Clear title
   - Detailed description
   - Priority level (high/medium/low)
   - Estimated timeline
3. Key Success Metrics
4. Risk Considerations

Format your response in a clear, professional manner suitable for a business document."""

        strategy = self._llm_generate(strategy_prompt, temperature=0.7)
        
        # Phase 3: Content Generation
        content_prompt = f"""Based on this business strategy:

{strategy[:1000]}

Create:
1. A professional LinkedIn post (150-200 words) that highlights the key strategy points
2. A brief executive summary (100 words) for internal communication

Format as:
LinkedIn Post:
[content]

Executive Summary:
[content]"""

        content_output = self._llm_generate(content_prompt, temperature=0.8)
        
        # Phase 4: Action Items Extraction
        action_items_prompt = f"""Extract specific, actionable items from this strategy:

{strategy[:1500]}

List 5-7 concrete action items in a numbered list format. Each item should be specific and actionable."""

        action_items = self._llm_generate(action_items_prompt, temperature=0.6)
        
        # Compile comprehensive result
        result = {
            "research_summary": industry_insights[:500] + ("..." if len(industry_insights) > 500 else ""),
            "strategy": strategy,
            "content_sample": content_output,
            "action_items": action_items,
            "sources": unique_results[:5],  # Limit to top 5 sources
            "industry": industry,
            "goal": goal,
            "generated_at": datetime.now().isoformat(),
            "confidence_score": min(len(unique_results) / 5.0, 1.0)  # Confidence based on research quality
        }
        
        logger.info(f"Pipeline completed successfully. Confidence: {result['confidence_score']:.2f}")
        return result

    def generate_doc(self, data: Dict, filename: str) -> str:
        """Enhanced document generation with better formatting"""
        try:
            doc = Document()
            
            # Title Page
            title = doc.add_heading('GenieSuite AI Strategy Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n").bold = True
            meta_para.add_run(f"Industry: {data.get('industry', 'N/A')}\n")
            meta_para.add_run(f"Business Goal: {data.get('goal', 'N/A')}")
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_text = data.get('research_summary', 'No summary available.')
            doc.add_paragraph(summary_text)
            
            # Strategy Section
            doc.add_heading('Strategic Plan', level=1)
            strategy_text = data.get('strategy', 'No strategy generated.')
            doc.add_paragraph(strategy_text)
            
            # Action Items
            if data.get('action_items'):
                doc.add_heading('Action Items', level=1)
                action_items_text = data.get('action_items', '')
                doc.add_paragraph(action_items_text)
            
            # Content Samples
            if data.get('content_sample'):
                doc.add_heading('Content Samples', level=1)
                content_text = data.get('content_sample', '')
                doc.add_paragraph(content_text)
            
            # Sources
            if data.get('sources'):
                doc.add_heading('Research Sources', level=1)
                for i, source in enumerate(data.get('sources', [])[:5], 1):
                    source_para = doc.add_paragraph()
                    source_para.add_run(f"{i}. ").bold = True
                    source_para.add_run(f"{source.get('title', 'Untitled')}\n")
                    if source.get('url'):
                        source_para.add_run(f"URL: {source.get('url')}").italic = True
            
            # Footer
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.add_run("Generated by GenieSuite AI\n").bold = True
            footer_para.add_run("For questions or feedback, please contact your system administrator.")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            os.makedirs("outputs", exist_ok=True)
            path = os.path.join("outputs", filename)
            doc.save(path)
            
            logger.info(f"Document generated successfully: {path}")
            return path
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            raise Exception(f"Failed to generate document: {str(e)}")


class MemoryStore:
    """Simple JSON file-based memory store for conversational state."""
    def __init__(self, path: str = "agent_memory.json"):
        self.path = path
        # initialize file
        if not os.path.exists(self.path):
            with open(self.path, "w", encoding="utf-8") as f:
                json.dump({}, f)

    def _read(self) -> Dict[str, Any]:
        try:
            with open(self.path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _write(self, data: Dict[str, Any]):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)

    def get_session(self, session_id: str) -> Dict[str, Any]:
        data = self._read()
        return data.get(session_id, {"history": []})

    def append_message(self, session_id: str, role: str, content: str):
        data = self._read()
        session = data.setdefault(session_id, {"history": []})
        session["history"].append({"role": role, "content": content, "ts": datetime.now().isoformat()})
        self._write(data)

    def clear_session(self, session_id: str):
        data = self._read()
        if session_id in data:
            del data[session_id]
            self._write(data)


class Agent:
    """Higher level agent built on top of GenieEngine that supports conversational patterns,
    intent parsing, action handlers and memory.
    """
    def __init__(self, engine: GenieEngine, memory_path: str = "agent_memory.json"):
        self.engine = engine
        self.memory = MemoryStore(memory_path)

    def parse_intent(self, message: str) -> Dict[str, Any]:
        """Basic intent parser. Uses keyword matching and falls back to LLM if available."""
        text = message.lower()
        # simple keyword intents
        if any(k in text for k in ["summarize", "summary", "summarise", "sum up"]):
            return {"intent": "summarize"}
        if any(k in text for k in ["plan", "strategy", "roadmap"]):
            return {"intent": "plan"}
        if any(k in text for k in ["todo", "action items", "tasks"]):
            return {"intent": "todo"}
        if any(k in text for k in ["search", "research", "find"]):
            return {"intent": "search"}
        if any(k in text for k in ["generate doc", "export", "document"]):
            return {"intent": "generate_doc"}

        # fallback: ask LLM for intent (if available)
        if self.engine.llm:
            prompt = f"Classify the following user message into one of: chat, summarize, plan, todo, search, generate_doc. Return only the intent.\nMessage:\n{message}"
            try:
                resp = self.engine._llm_generate(prompt, temperature=0.0)
                # normalize
                intent = resp.strip().split()[0].lower()
                return {"intent": intent}
            except Exception:
                pass

        return {"intent": "chat"}

    def handle_message(self, session_id: str, message: str) -> Dict[str, Any]:
        """Main conversational entrypoint. Appends to memory and routes to handlers."""
        # store user message
        self.memory.append_message(session_id, "user", message)

        intent_info = self.parse_intent(message)
        intent = intent_info.get("intent", "chat")

        if intent == "summarize":
            # extract text to summarize (simple heuristic)
            summary = self.summarize_text(message)
            self.memory.append_message(session_id, "agent", summary)
            return {"intent": "summarize", "response": summary}

        if intent == "plan":
            # expect the message contains a goal and optional industry
            # fallback: run pipeline with provided message as goal
            result = self.engine.run_pipeline(message, industry="General Business")
            out = result.get("strategy", "")
            self.memory.append_message(session_id, "agent", out)
            return {"intent": "plan", "response": out, "data": result}

        if intent == "todo":
            items = self.extract_action_items(message)
            self.memory.append_message(session_id, "agent", items)
            return {"intent": "todo", "response": items}

        if intent == "search":
            results = self.engine.search_web(message, max_results=5)
            summary = self.engine._extract_industry_insights("General", results)
            self.memory.append_message(session_id, "agent", summary)
            return {"intent": "search", "response": summary, "sources": results}

        if intent == "generate_doc":
            # generate a strategy and create a document
            result = self.engine.run_pipeline(message, industry="General Business")
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{os.urandom(4).hex()}.docx"
            path = self.engine.generate_doc(result, filename)
            resp = {"message": "Document generated", "path": path, "data": result}
            self.memory.append_message(session_id, "agent", "Generated document: " + filename)
            return {"intent": "generate_doc", "response": resp}

        # default: chat using LLM
        reply = self.chat_with_llm(session_id, message)
        self.memory.append_message(session_id, "agent", reply)
        return {"intent": "chat", "response": reply}

    def chat_with_llm(self, session_id: str, message: str) -> str:
        session = self.memory.get_session(session_id)
        history = session.get("history", [])[-10:]
        # build prompt
        prompt_parts = [f"{m['role']}: {m['content']}" for m in history]
        prompt_parts.append(f"user: {message}")
        prompt = "\n".join(prompt_parts)

        if self.engine.llm:
            try:
                resp = self.engine._llm_generate(prompt, temperature=0.7)
                return resp
            except Exception as e:
                logger.warning(f"LLM chat failed: {str(e)}")

        # fallback echo
        return f"I heard: {message}. (Configure an LLM for richer responses)"

    def summarize_text(self, text: str) -> str:
        prompt = f"Summarize the following text in 3-5 sentences:\n\n{text}"
        if self.engine.llm:
            try:
                return self.engine._llm_generate(prompt, temperature=0.3)
            except Exception:
                pass
        # fallback simple summary
        return (text[:800] + "...") if len(text) > 800 else text

    def extract_action_items(self, text: str) -> str:
        prompt = f"Extract 5 concrete action items from the following text:\n\n{text}"
        if self.engine.llm:
            try:
                return self.engine._llm_generate(prompt, temperature=0.4)
            except Exception:
                pass
        # fallback naive extraction: split by sentences
        sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 10]
        items = '\n'.join([f"{i+1}. {s[:200].strip()}" for i, s in enumerate(sentences[:5])])
        return items

